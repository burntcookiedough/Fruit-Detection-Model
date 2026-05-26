from __future__ import annotations

import json
from pathlib import Path
from datetime import date

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "doc"
ASSET_DIR = OUT_DIR / "assets"
OUT_DOCX = OUT_DIR / "fruit_detection_repo_analysis_report.docx"
STATS_PATH = ROOT / "tmp" / "repo_analysis_stats.json"

CLASSES = ["apple", "banana", "orange", "mango", "pineapple", "watermelon", "grapes", "pomegranate"]


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return ImageFont.truetype(c, size)
    return ImageFont.load_default()


def add_field(paragraph, instr: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instr
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr_text)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "244062")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_callout(doc: Document, title: str, body: str, fill="EAF2F8"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def draw_bar_chart(path: Path, title: str, counts: dict[str, int], width=1200, height=650):
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    f_title = font(34, True)
    f_label = font(22)
    f_small = font(20)
    d.text((50, 35), title, fill=(31, 78, 121), font=f_title)
    left, top, right, bottom = 250, 110, width - 80, height - 70
    max_v = max(counts.values()) or 1
    bar_h = (bottom - top) / len(counts) * 0.62
    palette = [(79, 129, 189), (155, 187, 89), (247, 150, 70), (128, 100, 162)]
    for idx, (name, value) in enumerate(counts.items()):
        y = top + idx * ((bottom - top) / len(counts)) + 8
        d.text((50, y + 6), name, fill=(40, 40, 40), font=f_label)
        bar_w = int((right - left) * value / max_v)
        color = palette[idx % len(palette)]
        d.rounded_rectangle((left, y, left + bar_w, y + bar_h), radius=6, fill=color)
        d.text((left + bar_w + 12, y + 4), f"{value:,}", fill=(40, 40, 40), font=f_small)
    d.line((left, bottom + 10, right, bottom + 10), fill=(180, 180, 180), width=2)
    img.save(path)


def draw_pipeline(path: Path):
    width, height = 1500, 700
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    title_f = font(36, True)
    box_f = font(22, True)
    small_f = font(18)
    d.text((50, 35), "Dataset and Model Pipeline", fill=(31, 78, 121), font=title_f)
    boxes = [
        ("Raw sources", "Kaggle, LVIS, Roboflow,\nripeness/quality YOLO sets", 70, 140),
        ("Canonical mapping", "Map labels into 8 fruit\nclasses, remove MD5 dupes", 370, 140),
        ("Quality gates", "Drop synthetic bias, loose\nboxes, bad banana samples", 670, 140),
        ("Leakage cleanup", "pHash/MD5 split leakage\nchecks and quarantine", 970, 140),
        ("Balanced V4 dataset", "17,876 train images;\nval/test left natural", 370, 390),
        ("YOLOv8s training", "640 px, batch 8, AMP,\ndisk cache, full aug", 670, 390),
        ("Evaluation", "Clean test plus synthetic\nwebcam stress holdout", 970, 390),
    ]
    for title, body, x, y in boxes:
        d.rounded_rectangle((x, y, x + 240, y + 125), radius=14, fill=(234, 242, 248), outline=(31, 78, 121), width=3)
        d.text((x + 16, y + 14), title, fill=(31, 78, 121), font=box_f)
        d.multiline_text((x + 16, y + 52), body, fill=(50, 50, 50), font=small_f, spacing=4)
    arrows = [
        ((310, 202), (370, 202)), ((610, 202), (670, 202)), ((910, 202), (970, 202)),
        ((790, 265), (520, 390)), ((610, 452), (670, 452)), ((910, 452), (970, 452)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        d.line((x1, y1, x2, y2), fill=(90, 90, 90), width=4)
        d.polygon([(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)], fill=(90, 90, 90))
    d.text((70, 570), "V5 direction: add class-targeted webcam-degraded training images while keeping validation/test clean.", fill=(90, 90, 90), font=font(22, True))
    img.save(path)


def draw_evolution(path: Path):
    width, height = 1500, 500
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((50, 35), "Project Evolution", fill=(31, 78, 121), font=font(36, True))
    stages = [
        ("V1", "Small synthetic/\nphotoshopped start"),
        ("V3", "Real merged data,\ncleaner train/test"),
        ("Phase A", "YOLOv8s baseline,\nmAP50 about 0.584"),
        ("V4", "Quality gates,\nleakage cleanup,\ntrain balancing"),
        ("V5", "Webcam-degraded\ntraining signal"),
    ]
    x0, y = 110, 210
    spacing = 300
    for i, (name, body) in enumerate(stages):
        x = x0 + i * spacing
        d.ellipse((x - 38, y - 38, x + 38, y + 38), fill=(79, 129, 189), outline=(31, 78, 121), width=3)
        d.text((x - 22, y - 17), name, fill="white", font=font(24, True))
        d.multiline_text((x - 95, y + 65), body, fill=(40, 40, 40), font=font(20), anchor=None, spacing=4)
        if i < len(stages) - 1:
            d.line((x + 45, y, x + spacing - 45, y), fill=(120, 120, 120), width=5)
            d.polygon([(x + spacing - 45, y), (x + spacing - 65, y - 11), (x + spacing - 65, y + 11)], fill=(120, 120, 120))
    img.save(path)


def section(doc: Document, num: int, title: str, level: int = 1):
    doc.add_heading(f"{num}. {title}", level=level)


def subsection(doc: Document, letter: str, title: str):
    doc.add_heading(f"{letter}. {title}", level=2)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    stats = json.loads(STATS_PATH.read_text(encoding="utf-8"))

    pipeline_png = ASSET_DIR / "pipeline.png"
    evolution_png = ASSET_DIR / "evolution.png"
    train_chart = ASSET_DIR / "train_distribution.png"
    test_chart = ASSET_DIR / "test_distribution.png"
    draw_pipeline(pipeline_png)
    draw_evolution(evolution_png)
    draw_bar_chart(train_chart, "V4 Training Box Distribution", stats["dataset_v4_balanced"]["train"]["class_counts"])
    draw_bar_chart(test_chart, "V4 Test Box Distribution", stats["dataset_v4_balanced"]["test"]["class_counts"])

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Segoe UI"
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Segoe UI"
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(68, 68, 68)

    for s in doc.sections:
        footer = s.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("Page ")
        add_field(footer, "PAGE")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Fruit Detection Model\nRepository Analysis Report")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Computer Vision / Deep Learning / YOLOv8 Engineering Review").italic = True
    doc.add_paragraph()
    meta = [
        ("Repository", str(ROOT)),
        ("Report date", date.today().isoformat()),
        ("Primary model family", "YOLOv8"),
        ("Current quality run", "runs/fruit_v4_quality"),
        ("Target classes", "8 fruit classes"),
    ]
    add_table(doc, ["Item", "Value"], meta, widths=[2.0, 4.8])
    add_callout(doc, "Verdict", "The repository is a practical, locally trainable fruit detector with unusually strong dataset-cleaning discipline for a small ML project. It is not production-ready yet because real webcam holdout evidence, automated tests, CI, and deployment packaging are still incomplete.", fill="D9EAD3")
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_paragraph("Note: In Microsoft Word, right-click the table above and choose Update Field if page numbers are not populated automatically.")
    doc.add_page_break()

    section(doc, 1, "Executive Summary")
    doc.add_paragraph("This repository implements an 8-class fruit object detector for local webcam and image inference. The technical center is a YOLOv8 training pipeline, with a major shift from early synthetic/limited datasets toward a cleaned, balanced, leakage-checked V4 dataset and a planned V5 webcam-robustness improvement.")
    add_bullets(doc, [
        "The strongest completed training run is V4 YOLOv8s: best validation mAP50 0.7543 and mAP50-95 0.5783 at epoch 93.",
        "A Phase A baseline reached only about 0.584 mAP50, so the V4 dataset and training changes produced a meaningful gain.",
        "The project is data-centric: raw-source auditing, class remapping, loose-box removal, pHash leakage checks, tiny-box filtering, and train-only balancing are all present.",
        "The main risk is deployment realism. V4 clean test performance is reasonable, but synthetic webcam stress results exposed severe apple/orange collapse, and the real webcam holdout is still empty.",
    ])

    section(doc, 2, "Repository Overview")
    add_table(doc, ["Area", "Important files/folders", "Purpose"], [
        ("Training", "train.py, config.py", "Fine-tune YOLOv8 models with centralized hyperparameters and resume support."),
        ("Evaluation", "evaluate.py, runs/detect/*", "Run Ultralytics validation on clean test or stress datasets and save plots."),
        ("Inference", "demo.py, inference/image.py, inference/webcam.py", "Run image and live webcam inference with bounding-box overlays."),
        ("Data pipeline", "pipeline/01_* through 09_*", "Audit, merge, clean, leakage-check, balance, and generate webcam-degraded data."),
        ("Datasets", "dataset_v4_balanced, dataset_v4_webcam_train, synthetic_webcam_holdout", "Final clean dataset, V5 training augmentation set, and stress holdout."),
        ("Exports", "export/export_onnx.py, export/export_tflite.py", "Deployment-oriented export entry points."),
        ("Models/runs", "models/best.pt, runs/fruit_v4_quality", "Champion weights and training artifacts."),
        ("Docs/dashboard", "README.md, FRUIT_V4_HANDOFF.md, dashboard.html", "Project summary, handoff, and local training dashboard."),
    ], widths=[1.5, 2.4, 3.2])
    doc.add_picture(str(pipeline_png), width=Inches(6.8))

    section(doc, 3, "Problem Statement")
    doc.add_paragraph("The project solves fruit localization and classification in images or webcam frames. It predicts bounding boxes and one of eight labels: apple, banana, orange, mango, pineapple, watermelon, grapes, and pomegranate.")
    add_bullets(doc, [
        "Useful applications include checkout assistance, inventory counting, education demos, food sorting prototypes, and edge-device computer vision experiments.",
        "The most realistic deployment scenario is a local laptop, Jetson-class edge device, or server-assisted camera feed. Mobile or very small CPU-only hardware would likely require a nano model, lower resolution, quantization, or hardware acceleration.",
        "The current scope is fruit type only. It does not detect ripeness, rot, bruising, or quality defects.",
    ])

    section(doc, 4, "Dataset Analysis")
    ds = stats["dataset_v4_balanced"]
    add_table(doc, ["Split", "Images", "Labels", "Boxes", "Empty labels"], [
        ("Train", f'{ds["train"]["images"]:,}', f'{ds["train"]["labels"]:,}', f'{ds["train"]["boxes"]:,}', ds["train"]["empty_labels"]),
        ("Validation", f'{ds["valid"]["images"]:,}', f'{ds["valid"]["labels"]:,}', f'{ds["valid"]["boxes"]:,}', ds["valid"]["empty_labels"]),
        ("Test", f'{ds["test"]["images"]:,}', f'{ds["test"]["labels"]:,}', f'{ds["test"]["boxes"]:,}', ds["test"]["empty_labels"]),
    ])
    add_table(doc, ["Class", "Train boxes", "Validation boxes", "Test boxes"], [
        (c, f'{ds["train"]["class_counts"][c]:,}', f'{ds["valid"]["class_counts"][c]:,}', f'{ds["test"]["class_counts"][c]:,}') for c in CLASSES
    ])
    doc.add_picture(str(train_chart), width=Inches(6.4))
    doc.add_picture(str(test_chart), width=Inches(6.4))
    add_callout(doc, "Dataset quality finding", "The V4 dataset is stronger than a typical student-level object detection dataset because it blocks training until labels, splits, duplicate leakage, empty labels, tiny boxes, and class coverage are checked. The remaining weakness is class imbalance in natural validation/test splits and limited real webcam evidence.", fill="EAF2F8")

    section(doc, 5, "Model Architecture")
    doc.add_paragraph("The completed quality model is YOLOv8s, a single-stage object detector from Ultralytics. It is trained from COCO-pretrained weights and adapted to eight fruit classes. The code also keeps YOLOv8n as a speed fallback.")
    add_table(doc, ["Component", "Implementation", "Engineering interpretation"], [
        ("Base weights", "yolov8s.pt", "Good local quality/speed balance on RTX 3060 6 GB."),
        ("Input", "640 x 640 images", "Reasonable for small fruit boxes without excessive VRAM pressure."),
        ("Output", "Bounding boxes, class labels, confidence", "Standard YOLO detection interface."),
        ("Losses", "YOLO box, classification, DFL losses", "Handled by Ultralytics; tracked in results.csv."),
        ("Model size", f'{stats["size_runs/fruit_v4_quality/weights/best.pt"]} MB', "Manageable for laptop and many edge accelerators."),
    ])

    section(doc, 6, "Training Pipeline")
    add_table(doc, ["Setting", "V4 value"], [
        ("Model", "YOLOv8s"),
        ("Dataset", "data_v4_balanced.yaml"),
        ("Epochs", "120"),
        ("Image size", "640"),
        ("Batch", "8"),
        ("Workers", "4"),
        ("Optimizer", "Ultralytics auto"),
        ("Schedule", "cos_lr=True"),
        ("AMP", "Enabled"),
        ("Cache", "disk"),
        ("Main augmentations", "HSV, rotation, translation, scale, shear, mosaic, mixup, erasing"),
    ])
    doc.add_paragraph("The pipeline is efficient for the available hardware because it uses disk caching rather than RAM caching, keeps batch size inside 6 GB VRAM, and preserves resume behavior. A minor issue is that Windows worker behavior can be fragile; the handoff correctly documents fallback to fewer workers if dataloader instability appears.")

    section(doc, 7, "Performance Evaluation")
    v4 = stats["run_fruit_v4_quality"]
    base = stats["run_fruit_v4_s_local"]
    add_table(doc, ["Run", "Epochs", "Best epoch", "Precision", "Recall", "mAP50", "mAP50-95"], [
        ("Phase A baseline", base["epochs_recorded"], base["best_mAP50_epoch"], f'{base["best_mAP50_precision"]:.3f}', f'{base["best_mAP50_recall"]:.3f}', f'{base["best_mAP50"]:.3f}', f'{base["best_mAP50_95"]:.3f}'),
        ("V4 YOLOv8s quality", v4["epochs_recorded"], v4["best_mAP50_epoch"], f'{v4["best_mAP50_precision"]:.3f}', f'{v4["best_mAP50_recall"]:.3f}', f'{v4["best_mAP50"]:.3f}', f'{v4["best_mAP50_95"]:.3f}'),
    ])
    add_table(doc, ["Class", "Clean test mAP50", "Synthetic webcam mAP50", "Interpretation"], [
        ("apple", "60.3%", "13.3%", "Critical webcam collapse."),
        ("banana", "56.8%", "30.2%", "Weak under degraded camera conditions."),
        ("orange", "64.4%", "10.7%", "Critical color-cast failure."),
        ("mango", "97.1%", "79.3%", "Strong and robust relative to others."),
        ("pineapple", "53.4%", "40.2%", "Moderate; needs more varied real samples."),
        ("watermelon", "51.5%", "59.0%", "Surprisingly robust in stress set."),
        ("grapes", "54.5%", "39.5%", "Moderate degradation."),
        ("pomegranate", "91.6%", "61.1%", "Good clean performance, still domain-sensitive."),
    ])
    for img in [ROOT / "runs/fruit_v4_quality/results.png", ROOT / "runs/fruit_v4_quality/confusion_matrix_normalized.png"]:
        if img.exists():
            doc.add_picture(str(img), width=Inches(6.5))
    add_callout(doc, "Performance interpretation", "The clean validation curve is credible, but deployment claims must stay conservative. Synthetic webcam degradation shows that color and image-quality shifts can break important classes, especially apple and orange.", fill="FCE5CD")

    section(doc, 8, "Code Quality Review")
    add_table(doc, ["Dimension", "Assessment"], [
        ("Organization", "Good separation between training, evaluation, inference, export, and data pipeline scripts."),
        ("Readability", "Generally clear names and module-level comments. Some docs are stale, for example README still references older v2/v3 paths in places."),
        ("Modularity", "Pipeline is script-based and understandable, but not packaged as reusable library modules."),
        ("Config management", "config.py centralizes core settings. This is practical, but V4/V5 switching through edits can cause accidental state drift."),
        ("Error handling", "Good checks for missing models/data and resume checkpoint. Dataset scripts use gates and SystemExit on failed quality thresholds."),
        ("Reproducibility", "Run folders, result CSVs, args.yaml, and handoff notes are strong. Missing formal experiment tracker and CI."),
        ("Testing", "No obvious automated unit/integration tests. Most validation is operational scripts and manual artifact inspection."),
    ])
    doc.add_paragraph("A useful code snippet from train.py is the resume branch: it loads last.pt and calls model.train(resume=True) without overriding hyperparameters. That is the right behavior for restoring optimizer, LR schedule, and epoch state.")

    section(doc, 9, "Deployment Readiness")
    add_table(doc, ["Deployment area", "Current state", "Readiness"], [
        ("Local webcam", "inference/webcam.py exists and overlays detections/FPS.", "Prototype-ready."),
        ("ONNX", "export/export_onnx.py exists.", "Ready to try, needs target-device benchmark."),
        ("TFLite", "export/export_tflite.py exists.", "Planned path, verify conversion/runtime."),
        ("Docker", "No Dockerfile found.", "Not ready."),
        ("API/backend", "No production API service found.", "Not ready."),
        ("Edge/mobile", "YOLOv8s may be acceptable on accelerators; YOLOv8n fallback documented.", "Feasible after benchmarking."),
        ("Real webcam validation", "webcam_holdout has zero images/labels.", "Blocked for final claims."),
    ])

    section(doc, 10, "Bottlenecks and Weaknesses")
    add_bullets(doc, [
        "Domain shift remains the biggest model risk. Synthetic webcam stress exposed color-cast and low-quality camera failures.",
        "Validation/test are natural and honest, but minority classes such as pineapple and watermelon still have far fewer boxes than apple/orange.",
        "V5 webcam training data currently has a small image/label count mismatch in the extracted file count: 6,856 images and 6,858 labels. That should be audited before full V5 training.",
        "The codebase depends heavily on scripts and manual commands. This is workable locally but weaker for team reproducibility.",
        "No CI, data schema tests, automated smoke test, model-card generation, or release packaging is present.",
        "The real webcam holdout is empty, so deployment quality is not proven.",
    ])

    section(doc, 11, "Improvement Recommendations")
    add_table(doc, ["Priority", "Recommendation", "Reason"], [
        ("High", "Build a real webcam_holdout with at least 25-50 labeled images per class.", "This is the only reliable way to validate deployment behavior."),
        ("High", "Audit dataset_v4_webcam_train image-label mismatch before V5 full training.", "Prevents silent missing/extra labels and training confusion."),
        ("High", "Add a small automated dataset validation test suite.", "Catch bad boxes, missing pairs, wrong class IDs, and leakage before training."),
        ("High", "Evaluate every candidate on the same clean test and webcam stress/real holdout.", "Avoids selecting models based on non-comparable runs."),
        ("Medium", "Export ONNX and benchmark latency on the actual target device.", "Deployment choice should be based on FPS and memory, not only mAP."),
        ("Medium", "Add experiment tracking or a structured CSV/JSON registry.", "Makes run comparisons and rollback easier."),
        ("Nice-to-have", "Add Docker or a locked environment file.", "Improves reproducibility across machines."),
        ("Nice-to-have", "Try YOLOv8n/YOLO11n fallback after V5 quality run.", "May offer better speed for edge/mobile scenarios."),
    ], widths=[1.1, 3.0, 3.0])

    section(doc, 12, "Final Technical Verdict")
    doc.add_paragraph("The project is technically solid for a local ML engineering prototype and stronger than many small object-detection repositories because it treats data quality as a first-class problem. The V4 model is usable for demonstrations and controlled image inference. It is not yet production-ready because the final validation gap is real camera data, not another training epoch.")
    add_table(doc, ["Dimension", "Verdict"], [
        ("Engineering maturity", "Intermediate: practical scripts, good handoff, weak automation/CI."),
        ("Research suitability", "Good as a comparative applied CV project."),
        ("Production suitability", "Not yet; requires real holdout, packaging, monitoring, and deployment tests."),
        ("Edge/mobile suitability", "Feasible, but target-device benchmarking and possible nano/quantized export are required."),
        ("Scalability", "Dataset process scales reasonably, but script orchestration and disk cache size will become constraints."),
    ])

    section(doc, 13, "Project Evolution and Iterative Development")
    doc.add_picture(str(evolution_png), width=Inches(6.8))
    subsection(doc, "A", "Development Journey")
    doc.add_paragraph("The development path moved from a lightweight detector trained on smaller and partly synthetic sources to a stricter V4 pipeline. The important engineering decision was to stop treating model architecture as the first lever and instead rebuild the dataset until label quality, split integrity, and class coverage were acceptable.")
    subsection(doc, "B", "Dataset Evolution")
    doc.add_paragraph("The dataset evolved from broad raw downloads into a canonical 8-class YOLO dataset. The pipeline excluded Fruit-360-style clean-background sources, mapped ripeness/quality labels back to fruit type where appropriate, skipped unsupported classes, removed exact duplicates, removed loose boxes, quarantined near-duplicate split leakage, filtered tiny boxes, and balanced train only.")
    subsection(doc, "C", "Iteration-by-Iteration Error Analysis")
    add_table(doc, ["Iteration", "Problem", "Root cause", "Fix attempt", "Result"], [
        ("Early synthetic/limited data", "Weak real-world generalization", "Clean backgrounds and narrow image variation", "Move to real merged sources", "Better domain coverage."),
        ("Phase A baseline", "mAP50 plateau around 0.584", "Dataset noise, imbalance, and leakage risk", "Build V4 quality pipeline", "V4 validation mAP50 reached about 0.754."),
        ("V4 clean model", "Apple/orange fail under webcam stress", "Color shortcuts and camera degradation shift", "V5 class-targeted webcam-degraded training images", "Planned/current next phase."),
        ("Resume attempt", "Accidental fresh run folder", "Used last.pt as model with new name", "Document correct --resume path", "Original run preserved and completed."),
    ])
    subsection(doc, "D", "Sub-Iterations and Experimental Changes")
    add_bullets(doc, [
        "Model size: YOLOv8s selected for quality; YOLOv8n retained as speed fallback.",
        "Input size: 640 chosen as the practical limit for small-object detail on 6 GB VRAM.",
        "Batch size: 8 used for YOLOv8s; smaller batches documented as OOM fallback.",
        "Augmentation: V4 used broad HSV/mosaic/mixup/erasing; V5 increases hue/saturation/value and webcam-specific degraded copies.",
        "DataLoader workers: workers 4 selected after cache-backed runs appeared stable, with workers 2 as fallback.",
    ])
    subsection(doc, "E", "Model Improvement Strategy")
    doc.add_paragraph("The improvement method is failure-driven: train, evaluate on clean and stress sets, inspect false positives/negatives, add or clean data for the observed failure mode, then compare on fixed splits. This is the right process. The missing maturity layer is automated experiment tracking and real holdout reporting.")
    subsection(doc, "F", "Resource Constraints and Hardware Limitations")
    doc.add_paragraph("The project is constrained by an RTX 3060 laptop GPU with 6 GB VRAM and 16 GB RAM. That explains the preference for YOLO nano/small models, batch 8 for YOLOv8s, disk caching instead of RAM caching, and avoiding larger model families until the data problem is solved. The documented tradeoff is appropriate: train a smaller clean model reliably rather than chasing a larger detector that slows iteration or fails with OOM.")
    subsection(doc, "G", "Engineering Tradeoff Analysis")
    add_table(doc, ["Tradeoff", "Decision", "Assessment"], [
        ("Accuracy vs speed", "YOLOv8s as quality model, YOLOv8n fallback", "Reasonable for laptop and edge exploration."),
        ("Dataset quality vs annotation effort", "Strict gates before training", "Correct; prevents wasted GPU time."),
        ("Synthetic stress vs real validation", "Synthetic holdout used temporarily", "Useful but not a replacement for real webcam labels."),
        ("Training time vs experimentation", "120 epoch quality runs with saved checkpoints", "Acceptable, but V5 will be slower due to extra data."),
        ("Complexity vs maintainability", "Script pipeline instead of full package", "Fast for one engineer, less robust for team production."),
    ])

    doc.add_heading("Appendix A. Evidence and Artifacts Reviewed", level=1)
    add_bullets(doc, [
        "README.md",
        "FRUIT_V4_HANDOFF.md",
        "MODEL_IMPROVEMENT_PIPELINE.md",
        "v5-webcam-training.md",
        "train.py, config.py, evaluate.py, demo.py",
        "pipeline/01_source_audit.py through pipeline/09_generate_webcam_train.py",
        "data_v4_balanced.yaml and data_v5_webcam.yaml",
        "runs/fruit_v4_quality/results.csv, args.yaml, plots, and weights",
        "Label statistics extracted from dataset_v4_balanced, dataset_v4_webcam_train, and synthetic_webcam_holdout",
    ])

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
